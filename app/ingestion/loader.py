# app/ingestion/loader.py

import os
from pathlib import Path
import json

# PDF (better)
import fitz  # pymupdf

# DOCX
from docx import Document as DocxDocument

# Langchain
from langchain_core.documents import Document


# ---------------------------
# Project paths
# ---------------------------
BASE_DIR = Path(__file__).resolve().parents[2]
DATA_PATH = BASE_DIR / "data" / "raw"          # ← this line was missing


# ---------------------------
# PDF Loader (STRUCTURE-AWARE)
# ---------------------------
def load_pdf(file_path: str):
    doc = fitz.open(file_path)
    pages = []

    for i, page in enumerate(doc):
        blocks = page.get_text("blocks")

        page_text = ""
        for b in blocks:
            text = b[4].strip()
            if text:
                page_text += text + "\n"

        pages.append({
            "text": page_text,
            "page": i
        })

    return pages


# ---------------------------
# File loader (single file)
# ---------------------------
def load_file(file_path: str):
    ext = Path(file_path).suffix.lower()

    # ---------------- TXT ----------------
    if ext == ".txt":
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()

    # ---------------- PDF ----------------
    elif ext == ".pdf":
        return load_pdf(file_path)

    # ---------------- DOCX ----------------
    elif ext == ".docx":
        doc = DocxDocument(file_path)
        return "\n".join([p.text for p in doc.paragraphs])

    # ---------------- JSON ----------------
    elif ext == ".json":
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return json.dumps(data, indent=2)

    else:
        raise ValueError(f"Unsupported file type: {ext}")


# ---------------------------
# Document loader (RAG-ready)
# ---------------------------
def load_documents(folder_path: str = str(DATA_PATH)):
    documents = []

    for root, _, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)

            try:
                content = load_file(file_path)

                # ---------------- PDF (multi-pages) ----------------
                if isinstance(content, list):
                    for page in content:
                        if page["text"].strip():
                            documents.append(
                                Document(
                                    page_content=page["text"],
                                    metadata={
                                        "source": file,
                                        "path": file_path,
                                        "page": page["page"]
                                    }
                                )
                            )

                # ---------------- TEXT / DOCX / JSON ----------------
                else:
                    if content.strip():
                        documents.append(
                            Document(
                                page_content=content,
                                metadata={
                                    "source": file,
                                    "path": file_path
                                }
                            )
                        )

            except Exception as e:
                print(f"[WARN] Failed to load {file_path}: {e}")

    if not documents:
        print("[INFO] No documents found in folder — starting with empty vector store.")
        return []

    return documents
